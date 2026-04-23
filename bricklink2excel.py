#!/usr/bin/env python3
"""
Bricklink XML to Excel/PDF Converter
Qt6 GUI — open a Bricklink wanted-list XML, preview with part images, export to Excel or PDF.
"""

import sys
import io
import json
import re
import tempfile
import xml.etree.ElementTree as ET
from pathlib import Path
from dataclasses import dataclass
from typing import Optional

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QTableWidget, QTableWidgetItem, QFileDialog,
    QLabel, QProgressBar, QHeaderView, QMessageBox, QStatusBar,
    QDialog, QDialogButtonBox, QCheckBox, QGroupBox,
    QLineEdit, QStyledItemDelegate,
)
from PyQt6.QtCore import Qt, QThread, pyqtSignal
from PyQt6.QtGui import QPixmap, QFont

import requests
from PIL import Image as PILImage

import openpyxl
from openpyxl.drawing.image import Image as XLImage
from openpyxl.drawing.spreadsheet_drawing import AnchorMarker, OneCellAnchor
from openpyxl.drawing.xdr import XDRPositiveSize2D
from openpyxl.styles import Font as XLFont, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors as rl_colors
from reportlab.lib.units import mm
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Image as RLImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

from docx import Document as DocxDocument
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------------------------------------------------------------------------
# Data
# ---------------------------------------------------------------------------

BRICKLINK_COLORS: dict[int, str] = {
    0: "White", 1: "Blue", 2: "Green", 3: "Dark Turquoise", 4: "Red",
    5: "Dark Pink", 6: "Brown", 7: "Light Gray", 8: "Dark Gray", 9: "Light Blue",
    10: "Bright Green", 11: "Black", 12: "Dark Orange", 14: "Yellow", 15: "Lime",
    17: "Sand Green", 18: "Flesh", 19: "Tan", 20: "Light Violet", 22: "Purple",
    23: "Dark Blue", 25: "Salmon", 27: "Dark Tan", 28: "Sand Blue", 30: "Light Pink",
    31: "Medium Orange", 36: "Medium Lime", 39: "Light Blue",
    41: "Trans-Dark Blue", 42: "Trans-Green", 43: "Trans-Medium Blue",
    44: "Trans-Red", 45: "Trans-Yellow", 46: "Trans-Orange", 47: "Trans-Smoke",
    52: "Trans-Purple", 54: "Trans-Light Blue", 57: "Trans-Pink",
    60: "Chrome Gold", 61: "Chrome Silver", 67: "Flat Silver",
    68: "Very Light Orange", 70: "Reddish Brown", 71: "Light Bluish Gray",
    72: "Dark Bluish Gray", 73: "Medium Blue", 74: "Medium Green",
    77: "Light Pink", 78: "Light Flesh", 79: "Milky White", 80: "Metallic Silver",
    82: "Metallic Gold", 84: "Very Light Brown", 85: "Dark Bluish Gray",
    86: "Light Bluish Gray", 87: "Dark Flesh", 88: "Dark Red", 89: "Dark Azure",
    90: "Medium Azure", 91: "Light Aqua", 93: "Olive Green", 94: "Medium Nougat",
    99: "Very Light Gray", 100: "Light Orange", 101: "Dark Brown", 110: "Violet",
    111: "Bright Pink", 115: "Medium Lime", 125: "Light Salmon", 128: "Dark Nougat",
    134: "Pearl Gold", 135: "Pearl Dark Gray", 137: "Metal Blue", 142: "Pearl Light Gold",
    143: "Trans-Very Light Blue", 148: "Pearl Dark Gray", 150: "Pearl Very Light Gray",
    151: "Very Light Bluish Gray", 176: "Red Metallic", 179: "Flat Silver",
    183: "Flat White", 191: "Bright Light Orange", 212: "Bright Light Blue",
    216: "Rust", 217: "Nougat", 222: "Coral", 223: "Pink",
    226: "Bright Light Yellow", 228: "Electric Blue", 229: "Dark Pink",
    230: "Lavender", 232: "Sky Blue", 272: "Dark Blue", 288: "Dark Green",
    294: "Glow in Dark Opaque", 297: "Pearl Gold", 308: "Dark Brown",
    313: "Maersk Blue", 320: "Dark Red", 321: "Dark Azure", 322: "Medium Azure",
    323: "Light Aqua", 324: "Pink", 325: "Lavender", 326: "Olive Green",
    334: "Chrome Gold", 335: "Sand Red", 366: "Earth Orange", 373: "Sand Purple",
    378: "Sand Green", 379: "Sand Blue", 380: "Sand Yellow", 383: "Chrome Silver",
    462: "Dark Orange", 484: "Dark Orange", 503: "Very Light Gray",
}

ITEM_TYPES = {
    "P": "Part", "S": "Set", "M": "Minifig", "B": "Book",
    "G": "Gear", "C": "Catalog", "I": "Instruction",
}

CONDITIONS = {"N": "New", "U": "Used", "X": "Any"}

THUMB_PX = 80   # pixel size of thumbnails shown in Qt table


# ---------------------------------------------------------------------------
# Domain model
# ---------------------------------------------------------------------------

@dataclass
class BricklinkItem:
    item_type: str
    item_id: str
    color_id: int
    max_price: float
    min_qty: int
    condition: str
    notify: str
    source_file: str = ""
    description: str = ""
    image_data: Optional[bytes] = None

    @property
    def color_name(self) -> str:
        return BRICKLINK_COLORS.get(self.color_id, f"Color #{self.color_id}")

    @property
    def type_label(self) -> str:
        return ITEM_TYPES.get(self.item_type, self.item_type)

    @property
    def condition_label(self) -> str:
        return CONDITIONS.get(self.condition, self.condition)

    @property
    def image_urls(self) -> list[str]:
        """Ordered list of URLs to try; first valid image response wins."""
        iid = self.item_id
        if self.item_type == "S":
            return [
                f"https://img.bricklink.com/ItemImage/ST/0/{iid}.t2.png",
                f"https://img.bricklink.com/ItemImage/SN/0/{iid}-1.jpg",
            ]
        if self.item_type == "M":
            return [
                f"https://img.bricklink.com/ItemImage/ST/0/{iid}.t2.png",
                f"https://img.bricklink.com/ItemImage/MN/0/{iid}.png",
            ]
        if self.item_type == "G":
            return [f"https://img.bricklink.com/ItemImage/GN/0/{iid}.png"]
        return [f"https://img.bricklink.com/ItemImage/PN/{self.color_id}/{iid}.png"]

    @property
    def catalog_url(self) -> str:
        type_char = {"P": "P", "S": "S", "M": "M", "G": "G",
                     "B": "B", "C": "C", "I": "I"}.get(self.item_type, "P")
        return (f"https://www.bricklink.com/v2/catalog/catalogitem.page"
                f"?{type_char}={self.item_id}")

    @property
    def price_label(self) -> str:
        return f"${self.max_price:.4f}" if self.max_price >= 0 else "Any"


# ---------------------------------------------------------------------------
# XML parser
# ---------------------------------------------------------------------------

def parse_xml(path: str) -> list[BricklinkItem]:
    tree = ET.parse(path)
    root = tree.getroot()
    filename = Path(path).stem
    items: list[BricklinkItem] = []
    for el in root.findall("ITEM"):
        def get(tag: str, default: str = "") -> str:
            return (el.findtext(tag) or default).strip()
        items.append(BricklinkItem(
            item_type=get("ITEMTYPE", "P"),
            item_id=get("ITEMID"),
            color_id=int(get("COLOR", "0")),
            max_price=float(get("MAXPRICE", "-1")),
            min_qty=int(get("MINQTY", "1")),
            condition=get("CONDITION", "N"),
            notify=get("NOTIFY", "N"),
            source_file=filename,
        ))
    return items


# ---------------------------------------------------------------------------
# Background image downloader
# ---------------------------------------------------------------------------

DESC_CACHE_FILE = Path(__file__).parent / "descriptions.json"


def _load_desc_cache() -> dict[str, str]:
    if DESC_CACHE_FILE.exists():
        try:
            return json.loads(DESC_CACHE_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {}


def _save_desc_cache(cache: dict[str, str]) -> None:
    try:
        DESC_CACHE_FILE.write_text(json.dumps(cache, indent=2, ensure_ascii=False),
                                   encoding="utf-8")
    except Exception:
        pass


def _parse_description(html: str) -> str:
    """Extract item name from a Bricklink catalog page.

    BL page title format:  {Description} : {Type} {id} | BrickLink
    BL meta description:   ItemName: LEGO {Description}, ItemType: ...
    """
    # 1. meta name="description" — structured field, most reliable
    m = re.search(r'ItemName:\s*(?:LEGO\s+)?(.+?)\s*,\s*ItemType', html)
    if m:
        return m.group(1).strip()

    # 2. page <title> — take the part BEFORE " : {Type} {id}"
    m = re.search(r'<title>([^<]+)</title>', html, re.IGNORECASE)
    if m:
        text = re.sub(r'\s*\|.*$', '', m.group(1)).strip()   # strip "| BrickLink"
        if ' : ' in text:
            text = text.split(' : ', 1)[0].strip()           # take before " : "
        return text

    return ""


class ImageDownloadThread(QThread):
    image_ready = pyqtSignal(int, bytes)   # (index, raw image bytes or b"")
    desc_ready  = pyqtSignal(int, str)     # (index, description text)
    status_msg  = pyqtSignal(str)
    all_done    = pyqtSignal()

    _HEADERS = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}

    def __init__(self, items: list[BricklinkItem], cache_dir: Path,
                 desc_cache: dict[str, str]):
        super().__init__()
        self.items = items
        self.cache_dir = cache_dir
        self.desc_cache = desc_cache
        self._abort = False

    def abort(self) -> None:
        self._abort = True

    def run(self) -> None:
        for i, item in enumerate(self.items):
            if self._abort:
                break

            # --- image ---
            cache_key = f"{item.item_type}_{item.color_id}_{item.item_id}.png"
            cache_file = self.cache_dir / cache_key
            if cache_file.exists():
                self.image_ready.emit(i, cache_file.read_bytes())
            else:
                self.status_msg.emit(f"Downloading image {item.item_id} …")
                data = b""
                for url in item.image_urls:
                    try:
                        resp = requests.get(url, timeout=10, headers=self._HEADERS)
                        if (resp.status_code == 200 and
                                resp.headers.get("Content-Type", "").startswith("image/")):
                            data = resp.content
                            break
                    except Exception:
                        continue
                if data:
                    cache_file.write_bytes(data)
                self.image_ready.emit(i, data)

            # --- description ---
            desc_key = f"{item.item_type}_{item.item_id}"
            if desc_key in self.desc_cache:
                self.desc_ready.emit(i, self.desc_cache[desc_key])
            else:
                self.status_msg.emit(f"Fetching description {item.item_id} …")
                try:
                    resp = requests.get(item.catalog_url, timeout=10,
                                        headers=self._HEADERS)
                    desc = _parse_description(resp.text) if resp.status_code == 200 else ""
                except Exception:
                    desc = ""
                self.desc_cache[desc_key] = desc
                _save_desc_cache(self.desc_cache)
                self.desc_ready.emit(i, desc)

        self.all_done.emit()


# ---------------------------------------------------------------------------
# Export helpers
# ---------------------------------------------------------------------------

def _to_thumbnail(raw: bytes, px: int) -> Optional[PILImage.Image]:
    try:
        img = PILImage.open(io.BytesIO(raw)).convert("RGBA")
        img.thumbnail((px, px), PILImage.LANCZOS)
        return img
    except Exception:
        return None


# Column metadata used by both exporters
_LEFT_ALIGN_COLS = {"Source File", "Description", "Color Name"}

_EXCEL_COL_WIDTHS: dict[str, int] = {
    "#": 5, "Image": 15, "Source File": 28, "Part ID": 12,
    "Description": 40, "Color ID": 10, "Color Name": 28,
    "Type": 10, "Qty": 7, "Condition": 12, "Max Price": 14,
}

_PDF_COL_WIDTHS_MM: dict[str, float] = {
    "#": 6, "Image": 17, "Source File": 24, "Part ID": 14,
    "Description": 38, "Color ID": 10, "Color Name": 26,
    "Type": 11, "Qty": 8, "Condition": 14, "Max Price": 15,
}


def _item_values(item: "BricklinkItem", idx: int) -> dict[str, object]:
    return {
        "#": idx + 1,
        "Image": "",
        "Source File": item.source_file,
        "Part ID": item.item_id,
        "Description": item.description,
        "Color ID": item.color_id,
        "Color Name": item.color_name,
        "Type": item.type_label,
        "Qty": item.min_qty,
        "Condition": item.condition_label,
        "Max Price": item.price_label,
    }


def export_excel(items: list["BricklinkItem"], path: str, col_cfg: dict[str, bool], col_order: list[str] | None = None) -> None:
    ordered = col_order if col_order else COLUMNS
    visible = [c for c in ordered if c in COLUMNS and col_cfg.get(c, True)]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Bricklink Inventory"

    header_font = XLFont(name="Calibri", bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill("solid", fgColor="1F3864")
    center   = Alignment(horizontal="center", vertical="center", wrap_text=True)
    left_mid = Alignment(horizontal="left", vertical="center", wrap_text=True)
    thin = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )
    even_fill = PatternFill("solid", fgColor="DCE6F1")

    for col_idx, col_name in enumerate(visible, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = center
        cell.border = thin
        ws.column_dimensions[get_column_letter(col_idx)].width = _EXCEL_COL_WIDTHS[col_name]

    ws.row_dimensions[1].height = 28
    ws.freeze_panes = "A2"

    has_img_col = "Image" in visible
    img_col_letter = get_column_letter(visible.index("Image") + 1) if has_img_col else None
    ROW_H_PT = 72 if has_img_col else 18

    for idx, item in enumerate(items):
        row = idx + 2
        ws.row_dimensions[row].height = ROW_H_PT
        fill = even_fill if idx % 2 == 0 else PatternFill("solid", fgColor="FFFFFF")
        data_font = XLFont(name="Calibri", size=10)
        values = _item_values(item, idx)

        for col_idx, col_name in enumerate(visible, 1):
            align = left_mid if col_name in _LEFT_ALIGN_COLS else center
            c = ws.cell(row=row, column=col_idx, value=values[col_name])
            c.fill = fill
            c.alignment = align
            c.border = thin
            c.font = data_font

        if has_img_col and item.image_data:
            pil_img = _to_thumbnail(item.image_data, 80)
            if pil_img:
                buf = io.BytesIO()
                pil_img.save(buf, format="PNG")
                buf.seek(0)
                xl_img = XLImage(buf)
                img_col_idx = visible.index("Image")
                col_w_px = int(ws.column_dimensions[img_col_letter].width * 7 + 5)
                row_h_px = int(ROW_H_PT * 4 / 3)
                img_w, img_h = pil_img.size
                x_emu = max(0, (col_w_px - img_w) // 2) * 9525
                y_emu = max(0, (row_h_px - img_h) // 2) * 9525
                marker = AnchorMarker(col=img_col_idx, colOff=x_emu,
                                     row=row - 1, rowOff=y_emu)
                xl_img.anchor = OneCellAnchor(
                    _from=marker,
                    ext=XDRPositiveSize2D(cx=img_w * 9525, cy=img_h * 9525)
                )
                ws.add_image(xl_img)

    wb.save(path)


def export_pdf(items: list["BricklinkItem"], path: str, col_cfg: dict[str, bool], col_order: list[str] | None = None) -> None:
    ordered = col_order if col_order else COLUMNS
    visible = [c for c in ordered if c in COLUMNS and col_cfg.get(c, True)]
    IMG_MM = 20

    MARGIN_MM = 15
    available_w = (210 - 2 * MARGIN_MM) * mm
    raw_widths = {c: _PDF_COL_WIDTHS_MM[c] for c in visible}
    total_raw = sum(raw_widths.values())
    col_widths = [raw_widths[c] / total_raw * available_w for c in visible]

    # Paragraph styles — font/alignment live here, not in TableStyle
    ps_hdr = ParagraphStyle("hdr", fontSize=9, fontName="Helvetica-Bold",
                            alignment=TA_CENTER, leading=11,
                            textColor=rl_colors.white, spaceBefore=0, spaceAfter=0)
    ps_center = ParagraphStyle("ctr", fontSize=8, fontName="Helvetica",
                               alignment=TA_CENTER, leading=10, spaceBefore=0, spaceAfter=0)
    ps_left   = ParagraphStyle("lft", fontSize=8, fontName="Helvetica",
                               alignment=TA_LEFT,   leading=10, spaceBefore=0, spaceAfter=0)

    tmp_dir = Path(tempfile.mkdtemp(prefix="bl_pdf_"))
    header_row = [Paragraph(c, ps_hdr) for c in visible]
    table_data: list[list] = [header_row]
    tmp_files: list[Path] = []

    for idx, item in enumerate(items):
        values = _item_values(item, idx)
        row: list[object] = []
        for col_name in visible:
            if col_name == "Image":
                cell: object = ""
                if item.image_data:
                    pil_img = _to_thumbnail(item.image_data, 100)
                    if pil_img:
                        tf = tmp_dir / f"img_{idx}.png"
                        pil_img.save(tf, format="PNG")
                        tmp_files.append(tf)
                        cell = RLImage(str(tf), width=IMG_MM * mm, height=IMG_MM * mm,
                                       kind="proportional")
                row.append(cell)
            else:
                ps = ps_left if col_name in _LEFT_ALIGN_COLS else ps_center
                row.append(Paragraph(str(values[col_name]), ps))
        table_data.append(row)

    style: list = [
        ("BACKGROUND",    (0, 0), (-1, 0), rl_colors.HexColor("#1F3864")),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("GRID",          (0, 0), (-1, -1), 0.4, rl_colors.grey),
        ("TOPPADDING",    (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ("LEFTPADDING",   (0, 0), (-1, -1), 4),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 4),
    ]
    for i in range(1, len(table_data)):
        if i % 2 == 0:
            style.append(("BACKGROUND", (0, i), (-1, i), rl_colors.HexColor("#DCE6F1")))

    table = Table(table_data, colWidths=col_widths, repeatRows=1)
    table.setStyle(TableStyle(style))

    title_style = getSampleStyleSheet()["Title"]
    doc = SimpleDocTemplate(
        path,
        pagesize=A4,
        rightMargin=MARGIN_MM * mm, leftMargin=MARGIN_MM * mm,
        topMargin=15 * mm, bottomMargin=15 * mm,
    )
    doc.build([Paragraph("Bricklink Inventory", title_style), table])

    for f in tmp_files:
        try:
            f.unlink()
        except OSError:
            pass


# ---------------------------------------------------------------------------
# Word export
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_valign(cell, align: str = "center") -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def export_word(items: list["BricklinkItem"], path: str,
                col_cfg: dict[str, bool], col_order: list[str] | None = None) -> None:
    ordered = col_order if col_order else COLUMNS
    visible = [c for c in ordered if c in COLUMNS and col_cfg.get(c, True)]

    doc = DocxDocument()

    # Page margins 1.5 cm each side
    for section in doc.sections:
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)

    title = doc.add_heading("Bricklink Inventory", level=1)
    title.alignment = 1  # center

    table = doc.add_table(rows=1, cols=len(visible))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, col_name in enumerate(visible):
        cell = hdr_cells[i]
        _set_cell_bg(cell, "1F3864")
        _set_cell_valign(cell)
        p = cell.paragraphs[0]
        p.alignment = 1  # center
        run = p.add_run(col_name)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    # Set column widths proportional to _PDF_COL_WIDTHS_MM
    page_w_cm = 21.0 - 3.0  # A4 minus margins
    total_raw = sum(_PDF_COL_WIDTHS_MM[c] for c in visible)
    col_widths_cm = [_PDF_COL_WIDTHS_MM[c] / total_raw * page_w_cm for c in visible]
    for i, w in enumerate(col_widths_cm):
        for row in table.rows:
            row.cells[i].width = Cm(w)

    tmp_dir = Path(tempfile.mkdtemp(prefix="bl_word_"))
    tmp_files: list[Path] = []

    for idx, item in enumerate(items):
        row = table.add_row()
        bg = "DCE6F1" if idx % 2 == 0 else "FFFFFF"
        values = _item_values(item, idx)

        for i, col_name in enumerate(visible):
            cell = row.cells[i]
            _set_cell_bg(cell, bg)
            _set_cell_valign(cell)

            if col_name == "Image":
                if item.image_data:
                    pil_img = _to_thumbnail(item.image_data, 80)
                    if pil_img:
                        tf = tmp_dir / f"img_{idx}.png"
                        pil_img.save(tf, format="PNG")
                        tmp_files.append(tf)
                        p = cell.paragraphs[0]
                        p.alignment = 1  # center
                        col_w_in = col_widths_cm[i] / 2.54
                        img_w_in = min(pil_img.width / 96, col_w_in * 0.9)
                        p.add_run().add_picture(str(tf), width=Inches(img_w_in))
            else:
                p = cell.paragraphs[0]
                p.alignment = 1 if col_name not in _LEFT_ALIGN_COLS else 0
                run = p.add_run(str(values[col_name]))
                run.font.size = Pt(8)

    doc.save(path)

    for f in tmp_files:
        try:
            f.unlink()
        except OSError:
            pass


# ---------------------------------------------------------------------------
# Config file
# ---------------------------------------------------------------------------

COLUMNS = ["#", "Image", "Source File", "Part ID", "Description",
           "Color ID", "Color Name", "Type", "Qty", "Condition", "Max Price"]
COL_IMG = 1
COL_SRC = 2

CONFIG_FILE = Path(__file__).parent / "config.json"
DEFAULT_COLUMN_VIS: dict[str, bool] = {col: True for col in COLUMNS}


def _row_key(item: "BricklinkItem") -> str:
    return f"{item.source_file}|{item.item_type}|{item.item_id}|{item.color_id}"


def _apply_saved_order(items: list, row_order: list[str]) -> list:
    """Sort items to match the saved row order; unknown items go to the end."""
    if not row_order:
        return items
    key_map = {_row_key(item): item for item in items}
    ordered = [key_map.pop(k) for k in row_order if k in key_map]
    ordered.extend(key_map.values())   # items not in saved order appended at end
    return ordered


def load_config() -> tuple[dict[str, bool], list[str], list[str]]:
    """Returns (column_visibility, row_order, col_order)."""
    default_cols = dict(DEFAULT_COLUMN_VIS)
    if CONFIG_FILE.exists():
        try:
            data = json.loads(CONFIG_FILE.read_text(encoding="utf-8"))
            if "columns" in data:                             # new nested format
                cols = dict(DEFAULT_COLUMN_VIS)
                cols.update({k: bool(v) for k, v in data["columns"].items()
                             if k in DEFAULT_COLUMN_VIS})
                return cols, data.get("row_order", []), data.get("col_order", [])
            else:                                             # old flat format
                cols = dict(DEFAULT_COLUMN_VIS)
                cols.update({k: bool(v) for k, v in data.items()
                             if k in DEFAULT_COLUMN_VIS})
                return cols, [], []
        except Exception:
            pass
    return default_cols, [], []


def save_config(col_vis: dict[str, bool], row_order: list[str], col_order: list[str]) -> None:
    CONFIG_FILE.write_text(
        json.dumps({"columns": col_vis, "row_order": row_order, "col_order": col_order}, indent=2),
        encoding="utf-8"
    )


# ---------------------------------------------------------------------------
# Column config dialog
# ---------------------------------------------------------------------------

class ColumnConfigDialog(QDialog):
    def __init__(self, current: dict[str, bool], parent=None) -> None:
        super().__init__(parent)
        self.setWindowTitle("Show / Hide Columns")
        self.setMinimumWidth(260)

        self._checks: dict[str, QCheckBox] = {}

        group = QGroupBox("Visible columns")
        inner = QVBoxLayout(group)
        for col in COLUMNS:
            cb = QCheckBox(col)
            cb.setChecked(current.get(col, True))
            inner.addWidget(cb)
            self._checks[col] = cb

        buttons = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Ok | QDialogButtonBox.StandardButton.Cancel
        )
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)

        layout = QVBoxLayout(self)
        layout.addWidget(group)
        layout.addWidget(buttons)

    def result_config(self) -> dict[str, bool]:
        return {col: cb.isChecked() for col, cb in self._checks.items()}


# ---------------------------------------------------------------------------
# Qt6 Main Window
# ---------------------------------------------------------------------------

class DraggableTable(QTableWidget):
    """QTableWidget that properly moves entire rows on drag-drop."""
    rowOrderChanged = pyqtSignal()

    def dropEvent(self, event) -> None:
        if event.source() is not self:
            event.ignore()
            return

        drop_row = self.indexAt(event.position().toPoint()).row()
        if drop_row < 0:
            drop_row = self.rowCount()

        sel_rows = sorted(set(idx.row() for idx in self.selectedIndexes()))
        if not sel_rows:
            event.ignore()
            return

        drag_row = sel_rows[0]
        if drag_row == drop_row or drag_row + 1 == drop_row:
            event.ignore()
            return

        # takeItem removes items from the model without destroying them
        saved = [self.takeItem(drag_row, c) for c in range(self.columnCount())]
        self.removeRow(drag_row)          # destroys cell widgets (images), items are safe

        if drag_row < drop_row:
            drop_row -= 1

        self.insertRow(drop_row)
        for c, item in enumerate(saved):
            if item:
                self.setItem(drop_row, c, item)

        self.selectRow(drop_row)
        event.accept()
        self.rowOrderChanged.emit()       # triggers image refresh + order save


class ReadOnlyDelegate(QStyledItemDelegate):
    """Opens a read-only QLineEdit on double-click so text can be selected/copied."""
    def createEditor(self, parent, *_):
        editor = QLineEdit(parent)
        editor.setReadOnly(True)
        return editor

    def setEditorData(self, editor, index):
        editor.setText(index.data() or "")
        editor.selectAll()


class NumericItem(QTableWidgetItem):
    """QTableWidgetItem that sorts by numeric value instead of string."""
    def __lt__(self, other: QTableWidgetItem) -> bool:
        try:
            return float(self.text()) < float(other.text())
        except ValueError:
            return super().__lt__(other)


class MainWindow(QMainWindow):
    def __init__(self) -> None:
        super().__init__()
        self.setWindowTitle("Bricklink XML Viewer & Exporter")
        self.resize(1120, 720)

        self.items: list[BricklinkItem] = []
        self.cache_dir = Path(tempfile.mkdtemp(prefix="bl_cache_"))
        self._dl_thread: Optional[ImageDownloadThread] = None
        self._col_cfg, self._row_order, self._col_order = load_config()
        self._desc_cache = _load_desc_cache()

        self._build_ui()
        self._apply_column_visibility()
        self._apply_col_order()

    # ------------------------------------------------------------------
    # UI construction
    # ------------------------------------------------------------------

    def _build_ui(self) -> None:
        root = QWidget()
        self.setCentralWidget(root)
        vbox = QVBoxLayout(root)
        vbox.setContentsMargins(8, 8, 8, 8)
        vbox.setSpacing(6)

        # ---- toolbar ----
        hbox = QHBoxLayout()
        self.btn_open    = QPushButton("Open XML …")
        self.btn_excel   = QPushButton("Export Excel")
        self.btn_pdf     = QPushButton("Export PDF")
        self.btn_word    = QPushButton("Export Word")
        self.btn_columns = QPushButton("Columns …")
        for btn in (self.btn_open, self.btn_excel, self.btn_pdf, self.btn_word, self.btn_columns):
            btn.setMinimumHeight(34)
            btn.setFont(QFont("Segoe UI", 10))
            hbox.addWidget(btn)
        self.btn_excel.setEnabled(False)
        self.btn_pdf.setEnabled(False)
        self.btn_word.setEnabled(False)
        hbox.addStretch()
        self.lbl_file = QLabel("No file loaded")
        self.lbl_file.setFont(QFont("Segoe UI", 9))
        hbox.addWidget(self.lbl_file)
        vbox.addLayout(hbox)

        # ---- progress bar ----
        self.progress = QProgressBar()
        self.progress.setTextVisible(True)
        self.progress.setFormat("Downloading images … %v / %m")
        self.progress.setVisible(False)
        vbox.addWidget(self.progress)

        # ---- table ----
        self.table = DraggableTable()
        self.table.setColumnCount(len(COLUMNS))
        self.table.setHorizontalHeaderLabels(COLUMNS)
        self.table.setEditTriggers(QTableWidget.EditTrigger.DoubleClicked)
        self.table.setItemDelegate(ReadOnlyDelegate(self.table))
        self.table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.table.setAlternatingRowColors(True)
        self.table.verticalHeader().setDefaultSectionSize(92)
        self.table.horizontalHeader().setStretchLastSection(True)
        hh = self.table.horizontalHeader()
        hh.setSectionsMovable(True)
        hh.sectionMoved.connect(self._on_col_moved)
        hh.setSectionResizeMode(0, QHeaderView.ResizeMode.Fixed)
        hh.setSectionResizeMode(COL_IMG, QHeaderView.ResizeMode.Fixed)
        self.table.setColumnWidth(0, 40)
        self.table.setColumnWidth(COL_IMG, 100)
        self.table.setColumnWidth(COL_SRC, 200)
        self.table.setSortingEnabled(True)
        self.table.horizontalHeader().sortIndicatorChanged.connect(self._on_order_changed)
        self.table.rowOrderChanged.connect(self._on_order_changed)
        vbox.addWidget(self.table)

        # ---- status bar ----
        self.sb = QStatusBar()
        self.setStatusBar(self.sb)

        self.btn_open.clicked.connect(self._open_xml)
        self.btn_excel.clicked.connect(self._export_excel)
        self.btn_pdf.clicked.connect(self._export_pdf)
        self.btn_word.clicked.connect(self._export_word)
        self.btn_columns.clicked.connect(self._open_column_config)

    # ------------------------------------------------------------------
    # Column visibility
    # ------------------------------------------------------------------

    def _apply_column_visibility(self) -> None:
        for idx, col in enumerate(COLUMNS):
            self.table.setColumnHidden(idx, not self._col_cfg.get(col, True))

    def _apply_col_order(self) -> None:
        if not self._col_order:
            return
        hh = self.table.horizontalHeader()
        for visual_pos, col_name in enumerate(self._col_order):
            if col_name not in COLUMNS:
                continue
            logical = COLUMNS.index(col_name)
            current_visual = hh.visualIndex(logical)
            if current_visual != visual_pos:
                hh.moveSection(current_visual, visual_pos)

    def _on_col_moved(self, *_) -> None:
        hh = self.table.horizontalHeader()
        self._col_order = [COLUMNS[hh.logicalIndex(v)] for v in range(hh.count())]
        save_config(self._col_cfg, self._row_order, self._col_order)

    def _open_column_config(self) -> None:
        dlg = ColumnConfigDialog(self._col_cfg, parent=self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            self._col_cfg = dlg.result_config()
            save_config(self._col_cfg, self._row_order, self._col_order)
            self._apply_column_visibility()

    # ------------------------------------------------------------------
    # Slots
    # ------------------------------------------------------------------

    def _open_xml(self) -> None:
        paths, _ = QFileDialog.getOpenFileNames(
            self, "Open Bricklink XML Files", "", "XML Files (*.xml);;All Files (*)"
        )
        if not paths:
            return

        all_items: list[BricklinkItem] = []
        errors: list[str] = []
        for path in paths:
            try:
                all_items.extend(parse_xml(path))
            except Exception as exc:
                errors.append(f"{Path(path).name}: {exc}")

        if errors:
            QMessageBox.warning(self, "Parse Errors", "\n".join(errors))
        if not all_items:
            return

        self.items = _apply_saved_order(all_items, self._row_order)
        file_count = len(paths)
        names = ", ".join(Path(p).name for p in paths)
        self.lbl_file.setText(names if file_count == 1 else f"{file_count} files loaded")
        self.sb.showMessage(f"Loaded {len(self.items)} items from {file_count} file(s).")
        self._populate_table()
        self.btn_excel.setEnabled(True)
        self.btn_pdf.setEnabled(True)
        self.btn_word.setEnabled(True)
        self._start_downloads()

    def _export_excel(self) -> None:
        path, _ = QFileDialog.getSaveFileName(
            self, "Save Excel File", "", "Excel Files (*.xlsx)"
        )
        if not path:
            return
        if not path.lower().endswith(".xlsx"):
            path += ".xlsx"
        try:
            export_excel(self.items, path, self._col_cfg, self._col_order)
            QMessageBox.information(self, "Exported", f"Saved:\n{path}")
        except Exception as exc:
            QMessageBox.critical(self, "Export Error", str(exc))

    def _export_pdf(self) -> None:
        path, _ = QFileDialog.getSaveFileName(
            self, "Save PDF File", "", "PDF Files (*.pdf)"
        )
        if not path:
            return
        if not path.lower().endswith(".pdf"):
            path += ".pdf"
        try:
            export_pdf(self.items, path, self._col_cfg, self._col_order)
            QMessageBox.information(self, "Exported", f"Saved:\n{path}")
        except Exception as exc:
            QMessageBox.critical(self, "Export Error", str(exc))

    def _export_word(self) -> None:
        path, _ = QFileDialog.getSaveFileName(
            self, "Save Word File", "", "Word Files (*.docx)"
        )
        if not path:
            return
        if not path.lower().endswith(".docx"):
            path += ".docx"
        try:
            export_word(self.items, path, self._col_cfg, self._col_order)
            QMessageBox.information(self, "Exported", f"Saved:\n{path}")
        except Exception as exc:
            QMessageBox.critical(self, "Export Error", str(exc))

    # ------------------------------------------------------------------
    # Table population
    # ------------------------------------------------------------------

    def _populate_table(self) -> None:
        self.table.setSortingEnabled(False)
        self.table.setRowCount(len(self.items))
        for row, item in enumerate(self.items):
            self._set_row_text(row, item)
            self._set_image_widget(row, item)
        self.table.setSortingEnabled(True)

    def _set_row_text(self, row: int, item: BricklinkItem) -> None:
        def cell(txt: str, align=Qt.AlignmentFlag.AlignCenter) -> QTableWidgetItem:
            c = QTableWidgetItem(txt)
            c.setTextAlignment(align)
            return c

        def num(value: int | float, align=Qt.AlignmentFlag.AlignCenter) -> NumericItem:
            c = NumericItem(str(value))
            c.setTextAlignment(align)
            return c

        left = Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter

        seq = num(row + 1)
        seq.setData(Qt.ItemDataRole.UserRole, item)   # anchor for image re-mapping
        self.table.setItem(row, 0, seq)
        self.table.setItem(row, COL_IMG, QTableWidgetItem(""))  # sortable placeholder
        self.table.setItem(row, COL_SRC, cell(item.source_file, left))
        self.table.setItem(row, 3, cell(item.item_id))
        self.table.setItem(row, 4, cell(item.description, left))
        self.table.setItem(row, 5, num(item.color_id))
        self.table.setItem(row, 6, cell(item.color_name, left))
        self.table.setItem(row, 7, cell(item.type_label))
        self.table.setItem(row, 8, num(item.min_qty))
        self.table.setItem(row, 9, cell(item.condition_label))
        self.table.setItem(row, 10, cell(item.price_label))

    def _set_image_widget(self, row: int, item: BricklinkItem) -> None:
        lbl = QLabel()
        lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        if item.image_data:
            pm = QPixmap()
            pm.loadFromData(item.image_data)
            if not pm.isNull():
                pm = pm.scaled(THUMB_PX, THUMB_PX,
                               Qt.AspectRatioMode.KeepAspectRatio,
                               Qt.TransformationMode.SmoothTransformation)
                lbl.setPixmap(pm)
            else:
                lbl.setText("N/A")
        else:
            lbl.setText("…")
            lbl.setStyleSheet("color: #999;")
        self.table.setCellWidget(row, COL_IMG, lbl)

    def _on_order_changed(self, *_) -> None:
        self._refresh_image_widgets()
        self._save_row_order()

    def _save_row_order(self) -> None:
        order: list[str] = []
        for row in range(self.table.rowCount()):
            cell = self.table.item(row, 0)
            if cell is not None:
                bl_item = cell.data(Qt.ItemDataRole.UserRole)
                if isinstance(bl_item, BricklinkItem):
                    order.append(_row_key(bl_item))
        self._row_order = order
        save_config(self._col_cfg, self._row_order, self._col_order)

    def _refresh_image_widgets(self) -> None:
        for row in range(self.table.rowCount()):
            cell = self.table.item(row, 0)
            if cell is None:
                continue
            bl_item = cell.data(Qt.ItemDataRole.UserRole)
            if isinstance(bl_item, BricklinkItem):
                self._set_image_widget(row, bl_item)

    def _row_of(self, item: BricklinkItem) -> int:
        for row in range(self.table.rowCount()):
            cell = self.table.item(row, 0)
            if cell is not None and cell.data(Qt.ItemDataRole.UserRole) is item:
                return row
        return -1

    # ------------------------------------------------------------------
    # Image downloading
    # ------------------------------------------------------------------

    def _start_downloads(self) -> None:
        if self._dl_thread and self._dl_thread.isRunning():
            self._dl_thread.abort()
            self._dl_thread.wait()

        self.progress.setMaximum(len(self.items))
        self.progress.setValue(0)
        self.progress.setVisible(True)

        self._dl_thread = ImageDownloadThread(self.items, self.cache_dir, self._desc_cache)
        self._dl_thread.image_ready.connect(self._on_image_ready)
        self._dl_thread.desc_ready.connect(self._on_desc_ready)
        self._dl_thread.all_done.connect(self._on_all_done)
        self._dl_thread.status_msg.connect(self.sb.showMessage)
        self._dl_thread.start()

    def _on_desc_ready(self, idx: int, desc: str) -> None:
        self.items[idx].description = desc
        row = self._row_of(self.items[idx])
        if row >= 0:
            left = Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter
            c = QTableWidgetItem(desc)
            c.setTextAlignment(left)
            self.table.setItem(row, 4, c)

    def _on_image_ready(self, idx: int, data: bytes) -> None:
        bl_item = self.items[idx]
        bl_item.image_data = data or None
        self.progress.setValue(self.progress.value() + 1)
        row = self._row_of(bl_item)
        if row >= 0:
            self._set_image_widget(row, bl_item)

    def _on_all_done(self) -> None:
        self.progress.setVisible(False)
        downloaded = sum(1 for i in self.items if i.image_data)
        self.sb.showMessage(
            f"Ready — {len(self.items)} parts, {downloaded} images loaded.", 8000
        )

    def keyPressEvent(self, event) -> None:
        if (event.modifiers() == Qt.KeyboardModifier.ControlModifier
                and event.key() == Qt.Key.Key_C):
            ranges = self.table.selectedRanges()
            if ranges:
                r = ranges[0]
                rows = []
                for row in range(r.topRow(), r.bottomRow() + 1):
                    cells = []
                    for col in range(r.leftColumn(), r.rightColumn() + 1):
                        if not self.table.isColumnHidden(col):
                            item = self.table.item(row, col)
                            cells.append(item.text() if item else "")
                    rows.append("\t".join(cells))
                QApplication.clipboard().setText("\n".join(rows))
            return
        super().keyPressEvent(event)

    def closeEvent(self, event) -> None:
        if self._dl_thread and self._dl_thread.isRunning():
            self._dl_thread.abort()
            self._dl_thread.wait()
        super().closeEvent(event)


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    win = MainWindow()
    win.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
